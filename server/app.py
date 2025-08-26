from fastapi import FastAPI, UploadFile, File, Form, Body
from fastapi.middleware.cors import CORSMiddleware
import fitz  
import google.generativeai as genai
import os, json, re
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

last_results = {}   # store last AI results for export
last_document_text = ""
last_document_text2 = ""  

@app.post("/analyze")
async def analyze_document(
    file1: UploadFile = File(...),
    file2: UploadFile = File(None),
    language: str = Form("English")
):
    global last_document_text, last_document_text2, last_results

    doc1 = fitz.open(stream=await file1.read(), filetype="pdf")
    text1 = "".join([page.get_text() for page in doc1])
    last_document_text = text1

    text2 = ""
    if file2:
        doc2 = fitz.open(stream=await file2.read(), filetype="pdf")
        text2 = "".join([page.get_text() for page in doc2])
        last_document_text2 = text2

    model = genai.GenerativeModel("gemini-2.5-flash-lite")

    if file2:
        summary1_prompt = f"Summarize this legal document in 4–6 sentences in {language} . simplify the language and techincial words:\n\n{text1}"
        summary2_prompt = f"Summarize this legal document in 4–6 sentences in {language} simplify the language and techincial words :\n\n{text2}"

        comparison_prompt = f"""
Compare the following two legal documents and return the result in **valid JSON** only in {language}. 
Do not include any explanations outside JSON. Use this format:

{{
  "comparison": [
    {{"aspect": "Aspect name", "doc1": "what doc1 says", "doc2": "what doc2 says"}}
  ],
  "favorability": {{
    "doc1": <number between 0 and 100 indicating strength for doc1>,
    "doc2": <number between 0 and 100 indicating strength for doc2>
  }}
}}

Document 1:
{text1}

Document 2:
{text2}
"""
        summary1 = model.generate_content(summary1_prompt).text
        summary2 = model.generate_content(summary2_prompt).text
        comparison_raw = model.generate_content(comparison_prompt).text

        try:
            match = re.search(r"\{.*\}", comparison_raw, re.DOTALL)
            if match:
                comparison_data = json.loads(match.group(0))
            else:
                comparison_data = {"comparison": [], "favorability": {"doc1": 50, "doc2": 50}}
        except Exception:
            comparison_data = {"comparison": [], "favorability": {"doc1": 50, "doc2": 50}}

        last_results = {
            "summary1": summary1,
            "summary2": summary2,
            "comparison": comparison_data.get("comparison", []),
            "favorability": comparison_data.get("favorability", {"doc1": 50, "doc2": 50})
        }

        return {
            "main": "✅ Analysis complete: AI review finished.",
            **last_results,
            "timeline": None
        }

    else:
        summary_prompt = f"Summarize this legal document in 4–6 sentences in {language} . simplify the language and techincial words:\n\n{text1}"
        comparison_prompt = f"""
Analyze this legal document and return **valid JSON only**.
Highlight these aspects: Fairness, Risks, Missing Clauses.
Format:

{{
  "comparison": [
    {{"aspect": "Fairness", "doc": "..."}},
    {{"aspect": "Risks", "doc": "..."}},
    {{"aspect": "Missing Clauses", "doc": "..."}}
  ]
}}

Document:
{text1}
"""
        timeline_prompt = f"Just give timeline and important dates of the document in the specified format:\n\nDate:What happens (in one or two words)\nDate2:What happens\n\nJust return dates nothing else:\n\n{text1}"

        summary = model.generate_content(summary_prompt).text
        comparison_raw = model.generate_content(comparison_prompt).text
        timeline = model.generate_content(timeline_prompt).text

        try:
            match = re.search(r"\{.*\}", comparison_raw, re.DOTALL)
            if match:
                comparison_data = json.loads(match.group(0))
            else:
                comparison_data = {"comparison": []}
        except Exception:
            comparison_data = {"comparison": []}

        last_results = {
            "summary": summary,
            "comparison": comparison_data.get("comparison", []),
            "timeline": timeline
        }

        return {
            "main": "✅ Analysis complete: AI review finished.",
            **last_results
        }


@app.post("/whatif")
async def what_if(query: str = Form(...)):
    global last_document_text
    if not last_document_text:
        return {"error": "No document analyzed yet. Please upload and analyze a file first."}

    model = genai.GenerativeModel("gemini-2.5-flash-lite")
    whatif_prompt = f"Based on this legal document:\n\n{last_document_text}\n\nAnswer this hypothetical scenario: {query}"
    response = model.generate_content(whatif_prompt).text

    return {"response": response}

@app.post("/export")
async def export_results(format: str = Body("pdf")):
    if not last_results:
        return {"error": "No analysis to export!"}

    if format == "pdf":
        return await export_pdf()
    elif format == "word":
        return await export_word()
    else:
        return {"error": "Unsupported format"}
# ---------- NEW EXPORT ENDPOINTS ----------
@app.get("/export/pdf")
async def export_pdf():
    if not last_results:
        return {"error": "No analysis to export!"}

    file_path = "analysis_export.pdf"
    c = canvas.Canvas(file_path, pagesize=letter)
    y = 750
    c.setFont("Helvetica", 12)
    for key, value in last_results.items():
        c.drawString(50, y, f"{key}:")
        y -= 20
        if isinstance(value, list):
            for v in value:
                c.drawString(70, y, str(v))
                y -= 20
        else:
            c.drawString(70, y, str(value))
            y -= 40
    c.save()
    return FileResponse(file_path, filename="analysis_export.pdf", media_type="application/pdf")

@app.get("/export/word")
async def export_word():
    if not last_results:
        return {"error": "No analysis to export!"}

    file_path = "analysis_export.docx"
    doc = Document()
    doc.add_heading("LexVision AI Analysis", level=1)
    for key, value in last_results.items():
        doc.add_heading(key, level=2)
        if isinstance(value, list):
            for v in value:
                doc.add_paragraph(str(v))
        else:
            doc.add_paragraph(str(value))
    doc.save(file_path)
    return FileResponse(file_path, filename="analysis_export.docx", media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# Serve frontend
app.mount("/", StaticFiles(directory="static", html=True), name="static")
